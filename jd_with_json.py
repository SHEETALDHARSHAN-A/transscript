from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import json
from docx import Document
from io import BytesIO
import fitz  # PyMuPDF for PDF processing
from agents import Agent, InputGuardrail, GuardrailFunctionOutput, Runner
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import asyncio
import os
import logging
from datetime import datetime
import uuid
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Multilingual Job Description Extractor API",
    description="Extract and analyze job descriptions from documents in multiple languages",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure as needed for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Simplified Pydantic models without translation fields
class SkillRequirement(BaseModel):
    skill: str
    weightage: int = Field(..., ge=0, le=10, description="Skill importance from 0-10")

class JobDetails(BaseModel):
    job_title: str
    job_description: str
    skills_required: List[SkillRequirement]
    job_location: str
    work_from_home: bool
    min_experience: int = Field(..., ge=0)
    max_experience: int = Field(..., ge=0)
    detected_language: str = Field(..., description="Primary language detected in the document")
    salary_range: Optional[str] = Field(None, description="Salary information if available")
    company_name: Optional[str] = Field(None, description="Company name if mentioned")

class DocumentValidation(BaseModel):
    is_valid_job_document: bool
    reasoning: str
    detected_language: str = Field(..., description="Primary language of the document")
    confidence_score: float = Field(..., ge=0.0, le=1.0, description="Confidence in job document validation")

class LanguageDetection(BaseModel):
    primary_language: str
    language_code: str  # ISO 639-1 code
    confidence: float = Field(..., ge=0.0, le=1.0)
    contains_multiple_languages: bool

class ProcessingResponse(BaseModel):
    task_id: str
    status: str
    message: str

# Update the JobExtractionResult model to only include job details
class JobExtractionResult(BaseModel):
    job_title: str
    job_description: str
    skills_required: List[SkillRequirement]
    job_location: str
    work_from_home: bool
    min_experience: int = Field(..., ge=0)
    max_experience: int = Field(..., ge=0)
    salary_range: Optional[str] = None
    company_name: Optional[str] = None

class TextExtractionRequest(BaseModel):
    text: str
    document_type: Optional[str] = "text"

# In-memory storage for background tasks (use Redis or database in production)
task_storage: Dict[str, Dict[str, Any]] = {}

def extract_text_from_pdf(pdf_source) -> str:
    """
    Extracts text from a PDF file.
    :param pdf_source: Path to the PDF file or BytesIO object.
    :return: Extracted text as a string.
    """
    text = ""
    try:
        if isinstance(pdf_source, str):
            pdf_document = fitz.open(pdf_source)
        elif isinstance(pdf_source, BytesIO):
            pdf_document = fitz.open(stream=pdf_source, filetype="pdf")
        else:
            raise ValueError("Unsupported PDF source type.")

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
        pdf_document.close()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {str(e)}")
    return text

def extract_text_from_word(docx_source) -> str:
    """
    Extracts text from a Word document.
    :param docx_source: Path to the Word file or BytesIO object.
    :return: Extracted text as a string.
    """
    text = ""
    try:
        if isinstance(docx_source, str):  # File path
            docx_document = Document(docx_source)
        elif isinstance(docx_source, BytesIO):  # BytesIO stream
            docx_document = Document(docx_source)
        else:
            raise ValueError("Unsupported docx source type.")

        for para in docx_document.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text from Word document: {str(e)}")
    
    return text

# Language detection agent
language_detector_agent = Agent(
    name="Multilingual Language Detector",
    instructions="""
    You are an expert multilingual language detection system. Analyze the provided text to:
    
    1. Identify the primary language of the document
    2. Provide the ISO 639-1 language code
    3. Assess confidence in language detection
    4. Determine if multiple languages are present
    
    Supported languages include but not limited to:
    - English (en), Spanish (es), French (fr), German (de), Italian (it)
    - Portuguese (pt), Dutch (nl), Russian (ru), Chinese (zh), Japanese (ja)
    - Korean (ko), Arabic (ar), Hindi (hi), Thai (th), Vietnamese (vi)
    - Polish (pl), Czech (cs), Hungarian (hu), Romanian (ro), Greek (el)
    - Hebrew (he), Turkish (tr), Swedish (sv), Norwegian (no), Danish (da), Finnish (fi)
    
    Be accurate in your language detection and provide high confidence scores only when certain.
    """,
    output_type=LanguageDetection,
)

# Enhanced multilingual document validation agent
multilingual_document_validator_agent = Agent(
    name="Multilingual Document Validator",
    instructions="""
    You are an expert multilingual document validator with expertise in job descriptions across different languages and cultures.
    
    Analyze the provided text to determine if it contains a valid job description, regardless of language.
    
    A valid job description should contain AT LEAST ONE of the following:
    - A job title or position name (in any language)
    - Job responsibilities, duties, or requirements (in any language)
    - Required skills, qualifications, or experience (in any language)
    - Company information or hiring details (in any language)
    - Any work-related content that could be part of a job posting
    
    Be VERY INCLUSIVE in your validation. Even if the document is informal, incomplete, or mixed with other content,
    if it contains ANY job-related information, mark it as valid.
    
    You should be able to recognize job-related terminology in multiple languages.
    Provide reasonable confidence scores (0.5 or higher for any job-related content).
    Only reject documents that are clearly NOT job-related (personal letters, academic papers, etc.).
    """,
    output_type=DocumentValidation,
)

# Simplified multilingual job extraction agent - no translation required
multilingual_job_extractor_agent = Agent(
    name="Multilingual Job Details Extractor",
    instructions="""
    You are an expert multilingual HR analyst specializing in job description analysis across all languages and cultures.
    
    Extract comprehensive job details from the provided text in their original language - DO NOT translate anything:
    
    Instructions:
    1. Extract the exact job title/position name as it appears in the document
    2. Extract the complete job description exactly as written
    3. Identify the top 10 most important skills exactly as mentioned in the document
    4. Extract location exactly as stated in the document
    5. Determine if remote/hybrid work is mentioned
    6. Extract experience requirements
    7. Identify the primary language of the document
    8. Extract salary information if available (in original format/currency)
    9. Extract company name if mentioned
    
    Keep everything in the original language - no translation is needed.
    Be thorough and accurate in your extraction while preserving the original text format.
    """,
    output_type=JobDetails,
)

# Enhanced guardrail function for multilingual validation
async def multilingual_document_validation_guardrail(ctx, agent, input_data):
    """Validates if the document contains a valid job description in any language"""
    validation_input = str(input_data)
    
    try:
        # First detect language
        language_result = await Runner.run(language_detector_agent, validation_input, context=ctx.context)
        language_info = language_result.final_output_as(LanguageDetection)
        
        # Then validate document with language context
        validation_input_with_context = f"""
        Detected Language: {language_info.primary_language} ({language_info.language_code})
        Confidence: {language_info.confidence}
        Multiple Languages: {language_info.contains_multiple_languages}
        
        Document Text:
        {validation_input}
        """
        
        result = await Runner.run(multilingual_document_validator_agent, validation_input_with_context, context=ctx.context)
        final_output = result.final_output_as(DocumentValidation)
        
        # Add language info to the validation result
        final_output.detected_language = language_info.primary_language
        
        # More lenient validation - lower confidence threshold and better error handling
        is_valid = final_output.is_valid_job_document and final_output.confidence_score >= 0.5
        
        if not is_valid:
            logger.warning(f"Document validation failed: {final_output.reasoning}, Confidence: {final_output.confidence_score}")
        
        return GuardrailFunctionOutput(
            output_info=final_output,
            tripwire_triggered=not is_valid,
        )
    except Exception as e:
        logger.error(f"Error in validation guardrail: {e}")
        # More lenient approach - allow processing even if validation fails
        return GuardrailFunctionOutput(
            output_info=DocumentValidation(
                is_valid_job_document=True,  # Allow processing
                reasoning=f"Validation bypassed due to error: {str(e)}",
                detected_language="unknown",
                confidence_score=0.5
            ),
            tripwire_triggered=False,  # Don't trigger tripwire on error
        )

# Main multilingual job processing agent
multilingual_job_processing_agent = Agent(
    name="Multilingual Job Processing Agent",
    instructions="""
    You are a sophisticated multilingual job description processing system capable of handling documents in any language.
    
    Your capabilities include:
    1. Processing job descriptions in 50+ languages
    2. Extracting structured information in the original language
    3. Understanding cultural and regional differences in job descriptions
    4. Handling various document formats and layouts
    
    Process the document text and extract all relevant job details in the required structured format.
    Keep all extracted information in the original language - no translation required.
    
    Ensure high accuracy in extraction while preserving the original language and format.
    """,
    output_type=JobDetails,
    input_guardrails=[
        InputGuardrail(guardrail_function=multilingual_document_validation_guardrail),
    ],
)

# Direct extraction agent without guardrails for fallback
direct_job_extractor_agent = Agent(
    name="Direct Job Extractor",
    instructions="""
    You are a direct job information extractor. Extract job-related information from ANY text, even if it's not a perfect job description.
    
    Extract whatever job-related information you can find:
    - Any mention of job titles, positions, or roles
    - Any job responsibilities, tasks, or duties mentioned
    - Any skills, qualifications, or requirements
    - Location information if available
    - Experience requirements if mentioned
    - Company information if present
    - Salary information if available
    
    If specific information is not available, make reasonable assumptions or use "Not specified" values.
    For missing numeric values (experience), use 0 as default.
    
    Always try to extract something useful, even from incomplete or informal text.
    """,
    output_type=JobDetails,
)

async def extract_job_details_multilingual_async(document_text: str, document_type: str = "pdf") -> dict:
    """
    Async function to extract job details from multilingual documents using agents.
    No translation - keeps everything in original language.
    Uses fallback method if guardrails fail.
    """
    try:
        start_time = datetime.now()
        
        # Detect language first
        language_info = await detect_document_language_async(document_text)
        
        # Enhanced input formatting for original language extraction
        formatted_input = f"""
        Document Type: {document_type}
        
        MULTILINGUAL JOB DESCRIPTION EXTRACTION (ORIGINAL LANGUAGE ONLY)
        =================================================================

        Please extract comprehensive job details from the following document text.
        Keep all extracted information in the original language - no translation needed.
        
        Document Text:
        {document_text}
        
        Instructions:
        - Extract all information exactly as it appears in the document
        - Do not translate anything - keep original language
        - Maintain original formatting and terminology
        - Extract all available information in structured format
        """
        
        # Try with guardrails first
        try:
            logger.info("Attempting extraction with guardrails")
            result = await Runner.run(multilingual_job_processing_agent, formatted_input)
            
            if hasattr(result, 'final_output_as'):
                job_details = result.final_output_as(JobDetails).model_dump()
                processing_time = (datetime.now() - start_time).total_seconds()
                
                return {
                    "job_details": job_details,
                    "language_detection": language_info,
                    "processing_time": processing_time,
                    "status": "success",
                    "method": "guardrails"
                }
        except Exception as guardrail_error:
            logger.warning(f"Guardrail extraction failed: {guardrail_error}")
            
            # Fallback to direct extraction without guardrails
            try:
                logger.info("Attempting fallback extraction without guardrails")
                result = await Runner.run(direct_job_extractor_agent, formatted_input)
                
                if hasattr(result, 'final_output_as'):
                    job_details = result.final_output_as(JobDetails).model_dump()
                    processing_time = (datetime.now() - start_time).total_seconds()
                    
                    return {
                        "job_details": job_details,
                        "language_detection": language_info,
                        "processing_time": processing_time,
                        "status": "success",
                        "method": "direct_fallback",
                        "note": "Processed using fallback method due to guardrail validation"
                    }
                else:
                    return {
                        "job_details": result.final_output if result.final_output else {},
                        "language_detection": language_info,
                        "status": "partial_success",
                        "method": "direct_fallback"
                    }
            except Exception as fallback_error:
                logger.error(f"Both extraction methods failed: {fallback_error}")
                return {
                    "error": f"All extraction methods failed. Guardrail error: {str(guardrail_error)}, Fallback error: {str(fallback_error)}",
                    "status": "error"
                }
            
    except Exception as e:
        logger.error(f"Error extracting multilingual job details: {e}")
        return {
            "error": str(e),
            "status": "error"
        }

async def detect_document_language_async(document_text: str) -> dict:
    """
    Async function to detect the language of a document
    """
    try:
        result = await Runner.run(language_detector_agent, document_text)
        if hasattr(result, 'final_output_as'):
            language_info = result.final_output_as(LanguageDetection)
            return language_info.model_dump()
        else:
            return result.final_output if result.final_output else {}
    except Exception as e:
        logger.error(f"Error detecting language: {e}")
        return {"error": str(e)}

async def process_job_extraction_task(task_id: str, document_text: str, document_type: str):
    """Background task to process job extraction"""
    try:
        logger.info(f"Starting job extraction task {task_id}")
        
        # Update task status
        task_storage[task_id]["status"] = "processing"
        task_storage[task_id]["updated_at"] = datetime.now()
        
        # Process the document
        result = await extract_job_details_multilingual_async(document_text, document_type)
        
        # Update task with results
        task_storage[task_id].update({
            "status": "completed",
            "result": result,
            "completed_at": datetime.now(),
            "updated_at": datetime.now()
        })
        
        logger.info(f"Job extraction task {task_id} completed successfully")
        
    except Exception as e:
        logger.error(f"Error in job extraction task {task_id}: {e}")
        task_storage[task_id].update({
            "status": "failed",
            "error": str(e),
            "updated_at": datetime.now()
        })

# Add this function to handle JSON file saving
def save_job_details_to_file(job_details: dict, source_name: str) -> str:
    """
    Save job details to a JSON file.
    
    Args:
        job_details: Dictionary containing job details
        source_name: Name of the source file or "text_input"
    
    Returns:
        str: Path to the saved JSON file
    """
    # Create output directory if it doesn't exist
    output_dir = Path("d:/openai_agents/output")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Create filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"job_details_{source_name}_{timestamp}.json"
    output_path = output_dir / filename
    
    # Save JSON file with proper formatting and encoding
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(job_details, f, ensure_ascii=False, indent=2)
    
    return str(output_path)

# FastAPI Routes
@app.post("/extract/upload-sync")
async def extract_from_upload_sync(file: UploadFile = File(...)):
    """Upload and extract job details from PDF or Word document - returns results immediately"""
    try:
        # Validate file type
        allowed_types = ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]
        if file.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}. Supported types: PDF, DOCX, DOC")
        
        # Validate file size (optional - adjust as needed)
        max_file_size = 10 * 1024 * 1024  # 10MB
        content = await file.read()
        if len(content) > max_file_size:
            raise HTTPException(status_code=413, detail="File size too large. Maximum size is 10MB")
        
        # Extract text based on file type
        if file.content_type == "application/pdf":
            document_text = extract_text_from_pdf(BytesIO(content))
            document_type = "pdf"
        elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            document_text = extract_text_from_word(BytesIO(content))
            document_type = "docx"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        if not document_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the document. Please ensure the file contains readable text.")
        
        # Log the processing start
        logger.info(f"Processing file: {file.filename} ({file.content_type})")
        
        # Process the document synchronously
        result = await extract_job_details_multilingual_async(document_text, document_type)
        
        if result.get("status") == "error":
            raise HTTPException(status_code=500, detail=result.get("error"))
        
        # Get job details
        job_details = result.get("job_details", {})
        
        # Save to JSON file
        source_name = Path(file.filename).stem
        saved_path = save_job_details_to_file(job_details, source_name)
        
        # Return response with file path
        return {
            "job_details": job_details,
            "saved_to": saved_path
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file upload sync: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.post("/extract/text")
async def extract_from_text(request: TextExtractionRequest):
    """Extract job details from provided text - no translation"""
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text content cannot be empty")
        
        result = await extract_job_details_multilingual_async(request.text, request.document_type)
        
        if result.get("status") == "error":
            raise HTTPException(status_code=500, detail=result.get("error"))
        
        # Get job details
        job_details = result.get("job_details", {})
        
        # Save to JSON file
        saved_path = save_job_details_to_file(job_details, "text_input")
        
        # Return response with file path
        return {
            "job_details": job_details,
            "saved_to": saved_path
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing text extraction: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.post("/extract/direct", response_model=JobExtractionResult)
async def extract_direct_no_guardrails(request: TextExtractionRequest):
    """Extract job details directly without guardrails - for testing difficult documents"""
    try:
        if not request.text.strip():
            raise HTTPException(status_code=400, detail="Text content cannot be empty")
        
        logger.info(f"Processing direct extraction request (no guardrails)")
        
        start_time = datetime.now()
        
        # Detect language
        language_info = await detect_document_language_async(request.text)
        
        # Direct extraction without guardrails
        formatted_input = f"""
        Document Type: {request.document_type}
        
        DIRECT JOB EXTRACTION (NO VALIDATION)
        =====================================
        
        Extract job information from this text, even if it's not a perfect job description:
        
        {request.text}
        
        Extract whatever job-related information you can find and format it properly.
        """
        
        try:
            result = await Runner.run(direct_job_extractor_agent, formatted_input)
            
            if hasattr(result, 'final_output_as'):
                job_details = result.final_output_as(JobDetails).model_dump()
                processing_time = (datetime.now() - start_time).total_seconds()
                
                task_id = str(uuid.uuid4())
                
                return JobExtractionResult(
                    task_id=task_id,
                    status="success",
                    language_detection=LanguageDetection(**language_info) if language_info else None,
                    job_details=JobDetails(**job_details),
                    processing_time=processing_time,
                    timestamp=datetime.now()
                )
            else:
                raise Exception("No valid output from extraction agent")
                
        except Exception as e:
            logger.error(f"Error in direct extraction: {e}")
            return JobExtractionResult(
                task_id=str(uuid.uuid4()),
                status="error",
                error=f"Direct extraction failed: {str(e)}",
                timestamp=datetime.now()
            )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing direct extraction: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now()}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)